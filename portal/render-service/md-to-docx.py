#!/usr/bin/env python3
"""
md-to-docx.py — Markdown → .docx 轉換器（DOC-I6-02：品質升級版）
用法: python md-to-docx.py <input.md> <output.docx> [bpmn_assets_dir]

支援：
  - 乾淨文件範本：A4、1 吋邊界、CJK eastAsia 字型、標題層級（含 H2 底線）、表格樣式（表頭底色 + 斑馬列）
  - H1-H6 標題、段落、**粗體**、*斜體*、`行內程式碼`、[超連結](url)
  - 無序/有序清單（含縮排巢狀）
  - 程式碼區塊（``` ... ```）
  - <!-- PAGE_BREAK --> → docx 分頁符
  - 表格（GFM pipe table）：表頭底色、置中、斑馬列、Table Grid 框線
  - 圖片（依原始尺寸等比、最寬約頁寬；過小的裝飾圖略過）
  - ```bpmn blocks → 真 SVG 向量嵌入（Word 2016+，PNG 後備）；bpmn_assets_dir/<idx>.svg + <idx>.png
  - YAML front-matter 自動剝除（不顯示在內文）

匯入素材目錄（bpmn_assets_dir）每個 ```bpmn fence 依序對應：
  <idx>.svg  ← 前端 bpmn-js 渲染出的向量 SVG（首選，真嵌入）
  <idx>.png  ← 同一張圖的點陣後備（SVG 嵌入時當 fallback；無 svg 時直接用）
"""
import sys, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# A4 內容可用寬度（210mm - 2*1in 邊界 ≈ 6.27in）；圖片最寬留一點餘裕
CONTENT_WIDTH_IN = 6.1


# ── Clean document template ──────────────────────────────────────────────────

def _set_style_font(style, latin=None, eastasia=None, size_pt=None, bold=None, color=None):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    if latin:
        rfonts.set(qn('w:ascii'), latin)
        rfonts.set(qn('w:hAnsi'), latin)
        rfonts.set(qn('w:cs'), latin)
    if eastasia:
        rfonts.set(qn('w:eastAsia'), eastasia)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def _style_bottom_border(style, color='333333', sz='8'):
    pPr = style.element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), color)
    pbdr.append(bot)
    pPr.append(pbdr)


def _setup_template(doc):
    """套用乾淨範本：A4 版面 + CJK 字型 + 標題/內文樣式。"""
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)

    LATIN = 'Calibri'
    EA = 'Microsoft JhengHei'   # 微軟正黑體：乾淨現代的繁中預設

    # 內文
    normal = doc.styles['Normal']
    _set_style_font(normal, latin=LATIN, eastasia=EA, size_pt=11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.4

    # 標題層級（顏色近黑、對齊偏好對齊預覽）
    heads = {
        1: dict(size=20, color=RGBColor(0x11, 0x11, 0x11), center=True),
        2: dict(size=15, color=RGBColor(0x11, 0x11, 0x11), border=True),
        3: dict(size=13, color=RGBColor(0x1a, 0x1a, 0x1a)),
        4: dict(size=12, color=RGBColor(0x33, 0x33, 0x33)),
        5: dict(size=11, color=RGBColor(0x33, 0x33, 0x33)),
        6: dict(size=11, color=RGBColor(0x55, 0x55, 0x55)),
    }
    for lvl, cfg in heads.items():
        try:
            st = doc.styles['Heading %d' % lvl]
        except KeyError:
            continue
        _set_style_font(st, latin=LATIN, eastasia=EA, size_pt=cfg['size'],
                        bold=True, color=cfg['color'])
        st.paragraph_format.space_before = Pt(12 if lvl <= 2 else 9)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True
        if cfg.get('center'):
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cfg.get('border'):
            _style_bottom_border(st)


# ── Helper: hyperlink ────────────────────────────────────────────────────────

def _add_hyperlink(para, text, url):
    try:
        r_id = para.part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True)
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)
        hl.set(qn('w:history'), '1')
        rn = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        rn.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set(qn('xml:space'), 'preserve')
        rn.append(t)
        hl.append(rn)
        para._p.append(hl)
    except Exception:
        para.add_run(text)


# ── Helper: proportional image embed ─────────────────────────────────────────
# Returns True if embedded, False if skipped (tiny/decorative) or failed.

def _img_width(img_path, max_width_in, min_px=20):
    """回傳等比寬度（Inches）；過小的裝飾圖回傳 None（表示略過）。"""
    from PIL import Image as PILImage
    w_px = h_px = None
    dpi_x = 96.0
    try:
        with PILImage.open(img_path) as im:
            w_px, h_px = im.size
            dpi_info = im.info.get('dpi')
            if isinstance(dpi_info, (tuple, list)) and len(dpi_info) >= 1:
                try:
                    v = float(dpi_info[0])
                    if v > 0:
                        dpi_x = v
                except Exception:
                    pass
    except Exception:
        pass
    if w_px is not None and h_px is not None and (w_px < min_px or h_px < min_px):
        return None  # decorative: skip
    if w_px and dpi_x:
        actual_in = w_px / dpi_x
        return Inches(min(max_width_in, actual_in))
    return Inches(max_width_in)


def _embed_picture(para, img_path, max_width_in=CONTENT_WIDTH_IN, min_px=20):
    width = _img_width(img_path, max_width_in, min_px=min_px)
    if width is None:
        return False
    try:
        para.add_run().add_picture(img_path, width=width)
        return True
    except Exception:
        return False


# ── Helper: TRUE SVG embed (Word 2016+ vector, with PNG fallback) ─────────────
# Word stores a vector diagram as a normal <pic> whose <a:blip> references a PNG
# fallback PLUS an <asvg:svgBlip> extension referencing the real SVG part. Editors
# that understand SVG render the crisp vector; older ones show the PNG.

_A_SVG_NS = 'http://schemas.microsoft.com/office/drawing/2016/SVG/main'
_R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
_SVG_EXT_URI = '{96DAC541-7B7A-43D3-8B79-37D633B846F1}'


def _add_svg_part(document_part, svg_path):
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    with open(svg_path, 'rb') as f:
        svg_bytes = f.read()
    package = document_part.package
    partname = package.next_partname('/word/media/imagesvg%d.svg')
    svg_part = Part(partname, 'image/svg+xml', svg_bytes, package)
    return document_part.relate_to(svg_part, RT.IMAGE)


def _wire_svg_blip(inline_shape, svg_rid):
    from lxml import etree
    inline = inline_shape._inline
    blips = inline.xpath('.//a:blip')
    if not blips:
        return False
    blip = blips[0]
    extLst = OxmlElement('a:extLst')
    ext = OxmlElement('a:ext')
    ext.set('uri', _SVG_EXT_URI)
    svg_blip = etree.SubElement(ext, '{%s}svgBlip' % _A_SVG_NS)
    svg_blip.set('{%s}embed' % _R_NS, svg_rid)
    extLst.append(ext)
    blip.append(extLst)
    return True


def _embed_svg_with_fallback(doc, para, svg_path, png_path, max_width_in=CONTENT_WIDTH_IN):
    """以 png 為後備、svg 為向量主體嵌入；任何一步失敗就退回純 PNG。回傳是否成功。"""
    if not (png_path and os.path.isfile(png_path)):
        return False
    width = _img_width(png_path, max_width_in, min_px=1) or Inches(max_width_in)
    try:
        shape = para.add_run().add_picture(png_path, width=width)
    except Exception:
        return False
    # 試著掛上真 SVG；失敗就保留已嵌入的 PNG（不致更糟）
    if svg_path and os.path.isfile(svg_path):
        try:
            rid = _add_svg_part(doc.part, svg_path)
            _wire_svg_blip(shape, rid)
        except Exception:
            pass
    return True


# ── Helper: table cell shading ───────────────────────────────────────────────

def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


# ── Helper: inline markdown parser ──────────────────────────────────────────

_INLINE_RE = re.compile(
    r'(\*\*(?:[^*]|\*(?!\*))+\*\*)'    # **bold**
    r'|(__(?:[^_]|_(?!_))+__)'          # __bold__
    r'|(\*(?:[^*\n])+\*)'              # *italic*
    r'|(_(?:[^_\n])+_)'               # _italic_
    r'|(`[^`\n]+`)'                    # `code`
    r'|(!\[([^\]]*)\]\(([^)\n]+)\))'   # ![alt](src) — image first
    r'|(\[([^\]\n]+)\]\(([^)\n]+)\))'  # [text](url) — link
)


def _parse_inline(text, para, base_bold=False, base_italic=False, base_dir=None):
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            rn = para.add_run(text[pos:m.start()])
            rn.bold = base_bold
            rn.italic = base_italic

        tok = m.group(0)
        if tok.startswith('**') or tok.startswith('__'):
            rn = para.add_run(tok[2:-2])
            rn.bold = True
            rn.italic = base_italic
        elif tok.startswith('*') or tok.startswith('_'):
            rn = para.add_run(tok[1:-1])
            rn.italic = True
            rn.bold = base_bold
        elif tok.startswith('`'):
            rn = para.add_run(tok[1:-1])
            rn.font.name = 'Consolas'
            rn.font.size = Pt(10)
        elif tok.startswith('!['):
            alt = m.group(7) or ''
            src = m.group(8) or ''
            if base_dir:
                full = src if os.path.isabs(src) else os.path.normpath(os.path.join(base_dir, src))
                if os.path.isfile(full):
                    if _embed_picture(para, full):
                        pos = m.end()
                        continue
            rn = para.add_run('[圖片: %s]' % (alt or src))
            rn.italic = True
        elif tok.startswith('['):
            _add_hyperlink(para, m.group(10), m.group(11))

        pos = m.end()

    if pos < len(text):
        rn = para.add_run(text[pos:])
        rn.bold = base_bold
        rn.italic = base_italic


# ── Helper: horizontal rule ──────────────────────────────────────────────────

def _add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Helper: page break ───────────────────────────────────────────────────────

def _add_page_break(doc):
    p = doc.add_paragraph()
    rn = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    rn._r.append(br)


# ── Helper: strip YAML front-matter ──────────────────────────────────────────

def _strip_front_matter(content):
    return re.sub(r'^﻿?---\r?\n.*?\r?\n---[ \t]*\r?\n?', '', content, count=1, flags=re.S)


# ── Main converter ───────────────────────────────────────────────────────────

def convert(md_path, out_path, bpmn_dir=None):
    base_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, encoding='utf-8') as f:
        content = f.read()
    content = _strip_front_matter(content)

    doc = Document()
    _setup_template(doc)

    lines = content.split('\n')
    i = 0
    bpmn_idx = 0  # which ```bpmn fence we're on (matches frontend asset order)

    while i < len(lines):
        line = lines[i]
        s = line.rstrip()

        # ── Page break ───────────────────────────────────────────────
        if '<!-- PAGE_BREAK -->' in s.strip():
            _add_page_break(doc)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────
        if not s.strip():
            i += 1
            continue

        # ── Heading ──────────────────────────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=level)
            _parse_inline(m.group(2), h, base_dir=base_dir)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────
        if re.match(r'^[-*_]{3,}$', s.strip()):
            _add_hr(doc)
            i += 1
            continue

        # ── Fenced code block ────────────────────────────────────────
        if s.startswith('```'):
            lang = s[3:].strip().lower()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence

            if lang.startswith('bpmn'):
                svg_path = os.path.join(bpmn_dir, '%d.svg' % bpmn_idx) if bpmn_dir else None
                png_path = os.path.join(bpmn_dir, '%d.png' % bpmn_idx) if bpmn_dir else None
                embedded = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    embedded = _embed_svg_with_fallback(doc, p, svg_path, png_path, max_width_in=CONTENT_WIDTH_IN)
                except Exception:
                    embedded = False
                if not embedded:
                    # remove the empty centered paragraph, leave a clean placeholder
                    if p in doc.paragraphs:
                        p._element.getparent().remove(p._element)
                    ref = code_lines[0].strip() if code_lines else ''
                    p = doc.add_paragraph()
                    try:
                        p.style = doc.styles['Quote']
                    except Exception:
                        pass
                    rn = p.add_run('[BPMN 流程圖' + (': %s' % ref if ref else '') + ']')
                    rn.italic = True
                    rn.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                bpmn_idx += 1
            else:
                for cl in code_lines:
                    p = doc.add_paragraph()
                    rn = p.add_run(cl)
                    rn.font.name = 'Consolas'
                    rn.font.size = Pt(9.5)
            continue

        # ── Standalone image ─────────────────────────────────────────
        m_img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', s)
        if m_img:
            alt, src = m_img.group(1), m_img.group(2)
            full = src if os.path.isabs(src) else os.path.normpath(os.path.join(base_dir, src))
            embedded = False
            if os.path.isfile(full):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                embedded = _embed_picture(p, full)
                if not embedded and p in doc.paragraphs:
                    p._element.getparent().remove(p._element)
            if not embedded:
                p = doc.add_paragraph()
                rn = p.add_run('[圖片: %s]' % (alt or src))
                rn.italic = True
            i += 1
            continue

        # ── Unordered list ───────────────────────────────────────────
        m = re.match(r'^(\s*)[-*+]\s+(.*)', s)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            _parse_inline(m.group(2), p, base_dir=base_dir)
            i += 1
            continue

        # ── Ordered list ─────────────────────────────────────────────
        m = re.match(r'^(\s*)\d+[.)]\s+(.*)', s)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style='List Number')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            _parse_inline(m.group(2), p, base_dir=base_dir)
            i += 1
            continue

        # ── GFM table ────────────────────────────────────────────────
        if s.startswith('|') and s.endswith('|') and '|' in s[1:-1]:
            tbl_lines = [s]
            i += 1
            while i < len(lines):
                nl = lines[i].rstrip()
                if not (nl.startswith('|') and nl.endswith('|')):
                    break
                tbl_lines.append(nl)
                i += 1
            rows = [l for l in tbl_lines if not re.match(r'^[\|\s:=-]+$', l)]
            if rows:
                cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
                ncols = max(len(r) for r in cells)
                t = doc.add_table(rows=len(cells), cols=ncols)
                try:
                    t.style = doc.styles['Table Grid']
                except Exception:
                    pass
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ri, row in enumerate(cells):
                    for ci in range(ncols):
                        val = row[ci] if ci < len(row) else ''
                        cell = t.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        if ri == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _shade_cell(cell, 'D6E3F0')   # 表頭底色（對齊預覽）
                        elif ri % 2 == 1:
                            _shade_cell(cell, 'F5F5F5')    # 斑馬列
                        _parse_inline(val, p, base_bold=(ri == 0), base_dir=base_dir)
            continue

        # ── Regular paragraph (collect multi-line) ───────────────────
        para_parts = [s]
        i += 1
        while i < len(lines):
            nl = lines[i].rstrip()
            if (not nl.strip()
                    or nl.startswith('#')
                    or nl.startswith('```')
                    or nl.startswith('|')
                    or re.match(r'^\s*[-*+]\s+', nl)
                    or re.match(r'^\s*\d+[.)]\s+', nl)
                    or re.match(r'^[-*_]{3,}$', nl.strip())
                    or '<!-- PAGE_BREAK -->' in nl):
                break
            para_parts.append(nl)
            i += 1

        p = doc.add_paragraph()
        _parse_inline(' '.join(para_parts), p, base_dir=base_dir)

    doc.save(out_path)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python md-to-docx.py <input.md> <output.docx> [bpmn_assets_dir]', file=sys.stderr)
        sys.exit(1)
    bpmn_dir = sys.argv[3] if len(sys.argv) > 3 else None
    convert(sys.argv[1], sys.argv[2], bpmn_dir=bpmn_dir)
    print('OK: ' + sys.argv[2])
